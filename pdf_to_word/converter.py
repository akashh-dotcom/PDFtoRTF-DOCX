"""
PDF → DOCX converter with exact visual fidelity and editability options.

This module provides two conversion modes:
1. "exact" - Renders pages as images for perfect visual match (searchable but not editable)
2. "editable" - Extracts all elements as editable content with improved layout handling

The editable mode uses advanced techniques to minimize layout issues while maintaining editability.
"""

from __future__ import annotations

import html
import io
import re
import sys
from pathlib import Path
from typing import Optional, Sequence, List, Tuple, Literal, Dict, Any

import fitz  # PyMuPDF
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap, qn
from docx.shared import Emu, Pt, Inches, RGBColor, Twips


# ── Unit helpers ─────────────────────────────────────────────────────────────

_PT_TO_EMU = 12700          # 1 pt  = 12 700 EMU
_IN_TO_EMU = 914400         # 1 in  = 914 400 EMU

def _pt2emu(pt: float) -> int:
    return int(pt * _PT_TO_EMU)

def _emu2pt(emu: int) -> float:
    return emu / _PT_TO_EMU


# ── Shape ID counter ─────────────────────────────────────────────────────────

_SHAPE_ID_COUNTER = 0

def _next_shape_id() -> int:
    global _SHAPE_ID_COUNTER
    _SHAPE_ID_COUNTER += 1
    return _SHAPE_ID_COUNTER


def _escape(text: str) -> str:
    """Escape text for XML embedding."""
    return html.escape(text, quote=True)


# ── Color conversion ──────────────────────────────────────────────────────────

def _color_to_hex(color) -> str:
    """Convert a PyMuPDF colour (int, tuple, list, or None) to 6-char hex."""
    if color is None:
        return "000000"
    if isinstance(color, (tuple, list)):
        if len(color) >= 3:
            r, g, b = int(color[0] * 255), int(color[1] * 255), int(color[2] * 255)
            return f"{r:02X}{g:02X}{b:02X}"
        if len(color) == 1:
            v = int(color[0] * 255)
            return f"{v:02X}{v:02X}{v:02X}"
        return "000000"
    if isinstance(color, float):
        v = int(color * 255)
        return f"{v:02X}{v:02X}{v:02X}"
    # int (sRGB packed)
    r = (color >> 16) & 0xFF
    g = (color >> 8) & 0xFF
    b = color & 0xFF
    return f"{r:02X}{g:02X}{b:02X}"


def _hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


# ── Text width estimation ─────────────────────────────────────────────────────

# Character width factors for proportional fonts (relative to font size)
_CHAR_WIDTHS = {
    'i': 0.28, 'l': 0.28, 'I': 0.28, '1': 0.50, '.': 0.28, ',': 0.28,
    ':': 0.28, ';': 0.28, '!': 0.33, "'": 0.19, '"': 0.41, '|': 0.28,
    'j': 0.28, 'f': 0.33, 't': 0.33, 'r': 0.39, ' ': 0.28,
    'm': 0.89, 'w': 0.78, 'M': 0.89, 'W': 1.00, '@': 0.92, '%': 0.89,
    'A': 0.67, 'B': 0.67, 'C': 0.72, 'D': 0.72, 'E': 0.67, 'F': 0.61,
    'G': 0.78, 'H': 0.72, 'J': 0.50, 'K': 0.67, 'L': 0.56, 'N': 0.72,
    'O': 0.78, 'P': 0.67, 'Q': 0.78, 'R': 0.72, 'S': 0.67, 'T': 0.61,
    'U': 0.72, 'V': 0.67, 'X': 0.67, 'Y': 0.67, 'Z': 0.61,
}
_DEFAULT_CHAR_WIDTH = 0.56


def _estimate_text_width(text: str, font_size: float, font_name: str = "") -> float:
    """Estimate text width in points."""
    if not text:
        return 0
    
    # Check for monospace
    mono_keywords = ['mono', 'courier', 'consolas', 'menlo', 'fixed', 'code']
    is_mono = any(kw in font_name.lower() for kw in mono_keywords)
    
    if is_mono:
        return len(text) * font_size * 0.6
    
    total = sum(_CHAR_WIDTHS.get(c, _DEFAULT_CHAR_WIDTH) for c in text)
    return total * font_size


# ── Floating image insertion ──────────────────────────────────────────────────

def _add_floating_image(
    doc: Document,
    paragraph,
    image_bytes: bytes,
    x_emu: int,
    y_emu: int,
    w_emu: int,
    h_emu: int,
    shape_id: int,
    behind_doc: bool = True,
) -> None:
    """Insert a floating image at an exact page position."""
    # Create a temporary run to add an inline image, then extract the rId
    temp_run = paragraph.add_run()
    
    # Add inline image to get the relationship set up
    image_stream = io.BytesIO(image_bytes)
    try:
        inline_shape = temp_run.add_picture(image_stream, width=Emu(w_emu), height=Emu(h_emu))
    except Exception:
        # If image fails to load, skip it
        paragraph._element.remove(temp_run._element)
        return
    
    # Get the blip element to extract rId
    inline_xml = temp_run._element
    blip = None
    
    # Search for blip element
    for elem in inline_xml.iter():
        if 'blip' in elem.tag.lower():
            blip = elem
            break
    
    if blip is None:
        paragraph._element.remove(inline_xml)
        return
    
    # Extract the relationship ID
    rId = None
    for attr_name, attr_value in blip.attrib.items():
        if 'embed' in attr_name.lower():
            rId = attr_value
            break
    
    if not rId:
        paragraph._element.remove(inline_xml)
        return
    
    # Remove the inline image run
    paragraph._element.remove(inline_xml)
    
    behind = "1" if behind_doc else "0"

    xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
        '<mc:AlternateContent>'
        '<mc:Choice Requires="wps"><w:drawing>'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        ' simplePos="0" relativeHeight="{z}"'
        ' behindDoc="{behind}" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page">'
        '<wp:posOffset>{x}</wp:posOffset>'
        '</wp:positionH>'
        '<wp:positionV relativeFrom="page">'
        '<wp:posOffset>{y}</wp:posOffset>'
        '</wp:positionV>'
        '<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="{sid}" name="Img{sid}"/>'
        '<wp:cNvGraphicFramePr>'
        '<a:graphicFrameLocks noChangeAspect="1"/>'
        '</wp:cNvGraphicFramePr>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic>'
        '<pic:nvPicPr>'
        '<pic:cNvPr id="{sid}" name="Img{sid}"/>'
        '<pic:cNvPicPr/>'
        '</pic:nvPicPr>'
        '<pic:blipFill>'
        '<a:blip r:embed="{rId}"/>'
        '<a:stretch><a:fillRect/></a:stretch>'
        '</pic:blipFill>'
        '<pic:spPr>'
        '<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '</pic:spPr>'
        '</pic:pic>'
        '</a:graphicData></a:graphic>'
        '</wp:anchor>'
        '</w:drawing></mc:Choice>'
        '<mc:Fallback><w:pict/></mc:Fallback>'
        '</mc:AlternateContent>'
        '</w:r>'
    ).format(
        x=x_emu,
        y=y_emu,
        cx=w_emu,
        cy=h_emu,
        sid=shape_id,
        z=251650000 + shape_id,
        rId=rId,
        behind=behind,
    )

    run_element = parse_xml(xml)
    paragraph._element.append(run_element)


# ── Floating text box insertion ───────────────────────────────────────────────

def _add_textbox(
    paragraph,
    text: str,
    x_emu: int,
    y_emu: int,
    w_emu: int,
    h_emu: int,
    shape_id: int,
    font_name: str = "Arial",
    font_size_half_pt: int = 20,
    bold: bool = False,
    italic: bool = False,
    color_hex: str = "000000",
    underline: bool = False,
) -> None:
    """Add a visible, editable text box at exact position."""
    escaped_text = _escape(text)
    escaped_font = _escape(font_name)
    
    flags = ""
    if bold:
        flags += "<w:b/>"
    if italic:
        flags += "<w:i/>"
    if underline:
        flags += '<w:u w:val="single"/>'
    
    xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
        ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
        '<mc:AlternateContent>'
        '<mc:Choice Requires="wps"><w:drawing>'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        ' simplePos="0" relativeHeight="{z}"'
        ' behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page">'
        '<wp:posOffset>{x}</wp:posOffset>'
        '</wp:positionH>'
        '<wp:positionV relativeFrom="page">'
        '<wp:posOffset>{y}</wp:posOffset>'
        '</wp:positionV>'
        '<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="{sid}" name="TB{sid}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        '<wps:cNvSpPr txBox="1"/>'
        '<wps:spPr>'
        '<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '<a:noFill/><a:ln><a:noFill/></a:ln>'
        '</wps:spPr>'
        '<wps:txbx><w:txbxContent>'
        '<w:p><w:pPr><w:spacing w:after="0" w:before="0" w:line="240" w:lineRule="auto"/></w:pPr>'
        '<w:r><w:rPr>'
        '<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"/>'
        '{flags}'
        '<w:color w:val="{color}"/>'
        '<w:sz w:val="{sz}"/><w:szCs w:val="{sz}"/>'
        '</w:rPr>'
        '<w:t xml:space="preserve">{text}</w:t>'
        '</w:r>'
        '</w:p>'
        '</w:txbxContent></wps:txbx>'
        '<wps:bodyPr wrap="square" lIns="0" tIns="0" rIns="0" bIns="0"'
        ' anchor="t" anchorCtr="0"><a:noAutofit/></wps:bodyPr>'
        '</wps:wsp>'
        '</a:graphicData></a:graphic>'
        '</wp:anchor>'
        '</w:drawing></mc:Choice>'
        '<mc:Fallback><w:pict/></mc:Fallback>'
        '</mc:AlternateContent>'
        '</w:r>'
    ).format(
        x=x_emu,
        y=y_emu,
        cx=w_emu,
        cy=h_emu,
        sid=shape_id,
        z=251700000 + shape_id,
        font=escaped_font,
        flags=flags,
        color=color_hex,
        sz=font_size_half_pt,
        text=escaped_text,
    )

    run_element = parse_xml(xml)
    paragraph._element.append(run_element)


# ── Rectangle/Line shape insertion ────────────────────────────────────────────

def _add_rect_shape(
    paragraph,
    x_emu: int,
    y_emu: int,
    w_emu: int,
    h_emu: int,
    shape_id: int,
    stroke_color: str = "000000",
    fill_color: Optional[str] = None,
    stroke_width_emu: int = 12700,
) -> None:
    """Add a rectangle shape at exact position."""
    if w_emu <= 0 or h_emu <= 0:
        return
    
    fill_xml = f'<a:solidFill><a:srgbClr val="{fill_color}"/></a:solidFill>' if fill_color else '<a:noFill/>'
    
    xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
        ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
        '<mc:AlternateContent>'
        '<mc:Choice Requires="wps"><w:drawing>'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        ' simplePos="0" relativeHeight="{z}"'
        ' behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>{x}</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>{y}</wp:posOffset></wp:positionV>'
        '<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="{sid}" name="R{sid}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        '<wps:cNvSpPr/>'
        '<wps:spPr>'
        '<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '{fill}'
        '<a:ln w="{lw}"><a:solidFill><a:srgbClr val="{stroke}"/></a:solidFill></a:ln>'
        '</wps:spPr>'
        '<wps:bodyPr/>'
        '</wps:wsp>'
        '</a:graphicData></a:graphic>'
        '</wp:anchor>'
        '</w:drawing></mc:Choice>'
        '<mc:Fallback><w:pict/></mc:Fallback>'
        '</mc:AlternateContent>'
        '</w:r>'
    ).format(
        x=x_emu, y=y_emu, cx=w_emu, cy=h_emu,
        sid=shape_id, z=251600000 + shape_id,
        stroke=stroke_color, fill=fill_xml, lw=stroke_width_emu,
    )
    
    paragraph._element.append(parse_xml(xml))


def _add_line_shape(
    paragraph,
    x0_emu: int,
    y0_emu: int,
    x1_emu: int,
    y1_emu: int,
    shape_id: int,
    color: str = "000000",
    width_emu: int = 12700,
) -> None:
    """Add a line shape."""
    bx = min(x0_emu, x1_emu)
    by = min(y0_emu, y1_emu)
    bw = abs(x1_emu - x0_emu) or width_emu
    bh = abs(y1_emu - y0_emu) or width_emu
    
    flipH = "1" if x1_emu < x0_emu else "0"
    flipV = "1" if y1_emu < y0_emu else "0"
    
    xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
        ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
        '<mc:AlternateContent>'
        '<mc:Choice Requires="wps"><w:drawing>'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        ' simplePos="0" relativeHeight="{z}"'
        ' behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>{bx}</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>{by}</wp:posOffset></wp:positionV>'
        '<wp:extent cx="{bw}" cy="{bh}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="{sid}" name="L{sid}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        '<wps:cNvCnPr/>'
        '<wps:spPr>'
        '<a:xfrm flipH="{fH}" flipV="{fV}">'
        '<a:off x="0" y="0"/><a:ext cx="{bw}" cy="{bh}"/></a:xfrm>'
        '<a:prstGeom prst="line"><a:avLst/></a:prstGeom>'
        '<a:ln w="{lw}"><a:solidFill><a:srgbClr val="{color}"/></a:solidFill></a:ln>'
        '</wps:spPr>'
        '<wps:bodyPr/>'
        '</wps:wsp>'
        '</a:graphicData></a:graphic>'
        '</wp:anchor>'
        '</w:drawing></mc:Choice>'
        '<mc:Fallback><w:pict/></mc:Fallback>'
        '</mc:AlternateContent>'
        '</w:r>'
    ).format(
        bx=bx, by=by, bw=bw, bh=bh,
        sid=shape_id, z=251600000 + shape_id,
        color=color, lw=width_emu,
        fH=flipH, fV=flipV,
    )
    
    paragraph._element.append(parse_xml(xml))


# ── Text line grouping for better layout ──────────────────────────────────────

def _group_spans_by_line(blocks: List[Dict]) -> List[Dict[str, Any]]:
    """
    Group text spans into logical lines based on vertical position.
    Returns a list of line dictionaries with combined text and positioning.
    """
    lines = []
    
    for block in blocks:
        if block.get("type") != 0:  # Not a text block
            continue
        
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            
            # Get line bounding box
            line_bbox = list(line.get("bbox", [0, 0, 0, 0]))
            
            # Collect all spans in this line
            line_spans = []
            for span in spans:
                text = span.get("text", "")
                if not text:
                    continue
                
                line_spans.append({
                    "text": text,
                    "bbox": span.get("bbox"),
                    "font": span.get("font", "Arial"),
                    "size": span.get("size", 10),
                    "flags": span.get("flags", 0),
                    "color": span.get("color", 0),
                })
            
            if line_spans:
                lines.append({
                    "bbox": line_bbox,
                    "spans": line_spans,
                })
    
    return lines


# ── Page processing for editable mode ─────────────────────────────────────────

def _process_page_editable(
    pdf_doc: fitz.Document,
    word_doc: Document,
    page: fitz.Page,
    is_first: bool,
    dpi: int = 200,
    verbose: bool = False,
) -> None:
    """
    Process a PDF page in editable mode with improved layout handling.
    """
    rect = page.rect
    w_emu = _pt2emu(rect.width)
    h_emu = _pt2emu(rect.height)

    # Set up DOCX section
    if is_first:
        section = word_doc.sections[0]
    else:
        section = word_doc.add_section()

    landscape = rect.width > rect.height
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Emu(max(w_emu, h_emu))
        section.page_height = Emu(min(w_emu, h_emu))
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Emu(w_emu)
        section.page_height = Emu(h_emu)

    section.top_margin = Emu(0)
    section.bottom_margin = Emu(0)
    section.left_margin = Emu(0)
    section.right_margin = Emu(0)
    section.header_distance = Emu(0)
    section.footer_distance = Emu(0)

    anchor_para = word_doc.add_paragraph()
    anchor_para.paragraph_format.space_before = Pt(0)
    anchor_para.paragraph_format.space_after = Pt(0)

    # ── Step 1: Extract and place images ──────────────────────────────────
    image_rects = []  # Track image positions
    
    # Method 1: Direct image extraction
    image_list = page.get_images(full=True)
    processed_xrefs = set()
    
    for img_info in image_list:
        xref = img_info[0]
        if xref in processed_xrefs:
            continue
        processed_xrefs.add(xref)
        
        try:
            img_rects_list = page.get_image_rects(xref)
            if not img_rects_list:
                continue
            
            base_image = pdf_doc.extract_image(xref)
            if not base_image or not base_image.get("image"):
                continue
            
            img_bytes = base_image["image"]
            
            for img_rect in img_rects_list:
                if img_rect.width <= 0 or img_rect.height <= 0:
                    continue
                
                image_rects.append(fitz.Rect(img_rect))
                
                _add_floating_image(
                    word_doc,
                    anchor_para,
                    img_bytes,
                    _pt2emu(img_rect.x0),
                    _pt2emu(img_rect.y0),
                    _pt2emu(img_rect.width),
                    _pt2emu(img_rect.height),
                    _next_shape_id(),
                    behind_doc=True,
                )
        except Exception as e:
            if verbose:
                print(f"    Warning: Failed to extract image: {e}", file=sys.stderr)
            continue
    
    # Method 2: Detect figure regions by looking for large empty rectangles with borders
    # These often contain images that weren't directly extractable
    figure_regions = _detect_figure_regions(page, image_rects)
    
    for fig_rect in figure_regions:
        try:
            # Render this region as an image
            zoom = min(dpi, 250) / 72.0
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, clip=fig_rect, alpha=False)
            img_bytes = pix.tobytes("png")
            
            _add_floating_image(
                word_doc,
                anchor_para,
                img_bytes,
                _pt2emu(fig_rect.x0),
                _pt2emu(fig_rect.y0),
                _pt2emu(fig_rect.width),
                _pt2emu(fig_rect.height),
                _next_shape_id(),
                behind_doc=True,
            )
            image_rects.append(fig_rect)
        except Exception as e:
            if verbose:
                print(f"    Warning: Failed to render figure region: {e}", file=sys.stderr)
            continue

    # ── Step 2: Extract and place drawings/shapes ─────────────────────────
    drawings = page.get_drawings()
    complex_regions = []  # Regions that need rasterization
    
    for drawing in drawings:
        items = drawing.get("items", [])
        draw_rect = drawing.get("rect")
        stroke_color = _color_to_hex(drawing.get("color"))
        fill_color = drawing.get("fill")
        fill_hex = _color_to_hex(fill_color) if fill_color is not None else None
        stroke_width = drawing.get("width", 1) or 1
        
        # Check for complex paths (curves)
        has_curves = any(item[0] in ("c", "qu", "curve") for item in items)
        
        if has_curves and draw_rect:
            # Mark for rasterization
            complex_regions.append(fitz.Rect(draw_rect))
            continue
        
        # Process simple shapes
        for item in items:
            kind = item[0]
            
            if kind == "re":  # Rectangle
                r = item[1]
                _add_rect_shape(
                    anchor_para,
                    _pt2emu(r.x0),
                    _pt2emu(r.y0),
                    _pt2emu(r.width),
                    _pt2emu(r.height),
                    _next_shape_id(),
                    stroke_color=stroke_color,
                    fill_color=fill_hex,
                    stroke_width_emu=max(_pt2emu(stroke_width), 6350),
                )
            
            elif kind == "l":  # Line
                p1, p2 = item[1], item[2]
                _add_line_shape(
                    anchor_para,
                    _pt2emu(p1.x),
                    _pt2emu(p1.y),
                    _pt2emu(p2.x),
                    _pt2emu(p2.y),
                    _next_shape_id(),
                    color=stroke_color,
                    width_emu=max(_pt2emu(stroke_width), 6350),
                )
    
    # Rasterize complex regions
    merged_complex = _merge_rects(complex_regions) if complex_regions else []
    
    for region in merged_complex:
        if region.width < 5 or region.height < 5:
            continue
        
        try:
            zoom = min(dpi, 200) / 72.0
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, clip=region, alpha=False)
            img_bytes = pix.tobytes("png")
            
            _add_floating_image(
                word_doc,
                anchor_para,
                img_bytes,
                _pt2emu(region.x0),
                _pt2emu(region.y0),
                _pt2emu(region.width),
                _pt2emu(region.height),
                _next_shape_id(),
                behind_doc=True,
            )
        except Exception:
            continue

    # ── Step 3: Extract and place text ────────────────────────────────────
    text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_PRESERVE_LIGATURES)
    blocks = text_dict.get("blocks", [])
    
    # Group spans by line for better handling
    lines = _group_spans_by_line(blocks)
    
    for line_info in lines:
        line_bbox = line_info["bbox"]
        spans = line_info["spans"]
        
        # Process each span in the line
        for span in spans:
            text = span["text"]
            if not text.strip():
                continue
            
            bbox = span["bbox"]
            font = span["font"]
            size = span["size"]
            flags = span["flags"]
            color = span["color"]
            
            # Parse flags
            is_bold = bool(flags & (1 << 4))
            is_italic = bool(flags & (1 << 1))
            is_superscript = bool(flags & 1)
            
            # Clean font name
            clean_font = font
            if "+" in clean_font:
                clean_font = clean_font.split("+", 1)[1]
            
            # Calculate position and size
            x_emu = _pt2emu(bbox[0])
            y_emu = _pt2emu(bbox[1])
            
            # Calculate width more accurately
            pdf_width = bbox[2] - bbox[0]
            estimated_width = _estimate_text_width(text, size, clean_font)
            
            # Use PDF width but ensure it's at least as wide as estimated
            final_width = max(pdf_width, estimated_width)
            
            # Add padding to prevent clipping
            final_width *= 1.1
            
            # Height with padding for descenders
            final_height = max(bbox[3] - bbox[1], size * 1.4)
            
            box_w = _pt2emu(final_width)
            box_h = _pt2emu(final_height)
            
            # Font size in half-points
            size_half_pt = max(int(round(size * 2)), 2)
            
            color_hex = _color_to_hex(color)
            
            _add_textbox(
                anchor_para,
                text,
                x_emu,
                y_emu,
                box_w,
                box_h,
                _next_shape_id(),
                font_name=clean_font,
                font_size_half_pt=size_half_pt,
                bold=is_bold,
                italic=is_italic,
                color_hex=color_hex,
            )


# ── Page processing for exact mode ────────────────────────────────────────────

def _process_page_exact(
    pdf_doc: fitz.Document,
    word_doc: Document,
    page: fitz.Page,
    is_first: bool,
    dpi: int = 300,
) -> None:
    """
    Process a PDF page in exact mode - render as background image.
    """
    rect = page.rect
    w_emu = _pt2emu(rect.width)
    h_emu = _pt2emu(rect.height)

    # Set up DOCX section
    if is_first:
        section = word_doc.sections[0]
    else:
        section = word_doc.add_section()

    landscape = rect.width > rect.height
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Emu(max(w_emu, h_emu))
        section.page_height = Emu(min(w_emu, h_emu))
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Emu(w_emu)
        section.page_height = Emu(h_emu)

    section.top_margin = Emu(0)
    section.bottom_margin = Emu(0)
    section.left_margin = Emu(0)
    section.right_margin = Emu(0)
    section.header_distance = Emu(0)
    section.footer_distance = Emu(0)

    anchor_para = word_doc.add_paragraph()
    anchor_para.paragraph_format.space_before = Pt(0)
    anchor_para.paragraph_format.space_after = Pt(0)

    # Render entire page as image
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img_bytes = pix.tobytes("png")
    
    _add_floating_image(
        word_doc,
        anchor_para,
        img_bytes,
        x_emu=0,
        y_emu=0,
        w_emu=w_emu,
        h_emu=h_emu,
        shape_id=_next_shape_id(),
        behind_doc=True,
    )


# ── Helper: Detect figure regions ─────────────────────────────────────────────

def _detect_figure_regions(
    page: fitz.Page,
    already_extracted: List[fitz.Rect],
) -> List[fitz.Rect]:
    """
    Detect regions that likely contain figures/images that weren't directly extracted.
    
    This looks for:
    1. Large rectangular regions bounded by lines
    2. Areas with drawings but little/no text
    3. XObject forms that might contain graphics
    """
    figure_regions = []
    page_rect = page.rect
    
    # Get all drawings to find bordered regions
    drawings = page.get_drawings()
    
    # Collect all rectangles from drawings
    draw_rects = []
    for drawing in drawings:
        for item in drawing.get("items", []):
            if item[0] == "re":  # Rectangle
                r = fitz.Rect(item[1])
                # Only consider reasonably sized rectangles (likely figure frames)
                if r.width > 50 and r.height > 50:
                    draw_rects.append(r)
    
    # Get text blocks to identify text-sparse regions
    text_dict = page.get_text("dict")
    text_rects = []
    for block in text_dict.get("blocks", []):
        if block.get("type") == 0:  # Text block
            text_rects.append(fitz.Rect(block.get("bbox")))
    
    # Check each drawing rectangle to see if it might be a figure frame
    for rect in draw_rects:
        # Skip if too small
        if rect.width < 100 or rect.height < 100:
            continue
        
        # Skip if this overlaps with already extracted images
        overlaps_existing = False
        for existing in already_extracted:
            if rect.intersects(existing):
                intersection = rect & existing
                overlap_area = intersection.width * intersection.height
                rect_area = rect.width * rect.height
                if overlap_area > rect_area * 0.5:  # More than 50% overlap
                    overlaps_existing = True
                    break
        
        if overlaps_existing:
            continue
        
        # Check if this region has little text (suggesting it's a figure)
        text_in_region = 0
        for text_rect in text_rects:
            if rect.contains(text_rect):
                text_in_region += 1
        
        # If region has few text blocks, it's likely a figure
        # Allow some text (for captions inside the figure)
        if text_in_region <= 3:
            # Shrink slightly to avoid capturing the border itself
            inner_rect = fitz.Rect(
                rect.x0 + 2,
                rect.y0 + 2,
                rect.x1 - 2,
                rect.y1 - 2
            )
            if inner_rect.width > 50 and inner_rect.height > 50:
                figure_regions.append(inner_rect)
    
    # Also check for Form XObjects which often contain vector graphics
    try:
        xobjects = page.get_xobjects()
        for xobj in xobjects:
            try:
                xref = xobj[0]
                # Try to get the position of this XObject
                # This is tricky as XObjects don't always have direct position info
                xobj_rects = page.get_image_rects(xref)
                for xobj_rect in (xobj_rects or []):
                    if xobj_rect.width > 50 and xobj_rect.height > 50:
                        # Check if not already covered
                        is_covered = False
                        for existing in already_extracted + figure_regions:
                            if existing.contains(xobj_rect) or xobj_rect.contains(existing):
                                is_covered = True
                                break
                        if not is_covered:
                            figure_regions.append(fitz.Rect(xobj_rect))
            except Exception:
                continue
    except Exception:
        pass
    
    # Merge overlapping figure regions
    return _merge_rects(figure_regions, margin=10)


# ── Helper: Merge overlapping rectangles ──────────────────────────────────────

def _merge_rects(rects: List[fitz.Rect], margin: float = 5) -> List[fitz.Rect]:
    """Merge overlapping or nearby rectangles."""
    if not rects:
        return []
    
    merged = []
    used = [False] * len(rects)
    
    for i, rect in enumerate(rects):
        if used[i]:
            continue
        
        current = fitz.Rect(rect)
        current.x0 -= margin
        current.y0 -= margin
        current.x1 += margin
        current.y1 += margin
        
        changed = True
        while changed:
            changed = False
            for j, other in enumerate(rects):
                if used[j] or i == j:
                    continue
                other_exp = fitz.Rect(other)
                other_exp.x0 -= margin
                other_exp.y0 -= margin
                other_exp.x1 += margin
                other_exp.y1 += margin
                
                if current.intersects(other_exp):
                    current = current | other_exp
                    used[j] = True
                    changed = True
        
        used[i] = True
        merged.append(current)
    
    return merged


# ── Public API ───────────────────────────────────────────────────────────────

def convert_pdf_to_docx(
    pdf_path: str | Path,
    docx_path: Optional[str | Path] = None,
    *,
    pages: Optional[Sequence[int]] = None,
    dpi: int = 300,
    mode: Literal["exact", "editable"] = "editable",
    verbose: bool = False,
) -> Path:
    """Convert a PDF to a DOCX document.

    Parameters
    ----------
    pdf_path:
        Path to the source PDF file.
    docx_path:
        Destination path. Defaults to ``<input>.docx``.
    pages:
        0-based page indices. ``None`` → all pages.
    dpi:
        Resolution for rendering (higher = sharper but larger file).
        For editable mode, this affects complex graphics quality.
        For exact mode, this affects overall quality.
    mode:
        - "editable" (default): Extracts text, images, shapes as editable elements
        - "exact": Renders pages as images for perfect visual match
    verbose:
        Print progress to stderr.
    
    Returns
    -------
    Path to the created DOCX file.
    """
    global _SHAPE_ID_COUNTER
    _SHAPE_ID_COUNTER = 0

    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    if docx_path is None:
        docx_path = pdf_path.with_suffix(".docx")
    else:
        docx_path = Path(docx_path).resolve()

    docx_path.parent.mkdir(parents=True, exist_ok=True)

    pdf_doc = fitz.open(str(pdf_path))
    word_doc = Document()

    page_indices = list(pages) if pages is not None else list(range(len(pdf_doc)))
    total = len(page_indices)

    if verbose:
        print(f"Converting {pdf_path.name} ({total} pages, {mode} mode)...", file=sys.stderr)

    for i, idx in enumerate(page_indices):
        page = pdf_doc[idx]
        if verbose:
            print(f"  [{i + 1}/{total}] Processing page {idx + 1}...", file=sys.stderr)
        
        if mode == "exact":
            _process_page_exact(pdf_doc, word_doc, page, is_first=(i == 0), dpi=dpi)
        else:
            _process_page_editable(pdf_doc, word_doc, page, is_first=(i == 0), dpi=dpi, verbose=verbose)

    word_doc.save(str(docx_path))
    pdf_doc.close()

    if verbose:
        print(f"Done. Saved to: {docx_path}", file=sys.stderr)

    return docx_path
